#!/usr/bin/env python3
"""直接抓取"中国历史"完整文章，生成 docx。"""

import json
import os
import re
import ssl
import urllib.request
import urllib.parse

try:
    from docx import Document
except ImportError:
    print("pip install python-docx")
    exit(1)

ctx = ssl.create_default_context()
ctx.check_hostname = False
ctx.verify_mode = ssl.CERT_NONE


def api_get(params):
    url = "https://zh.wikipedia.org/w/api.php?" + urllib.parse.urlencode(params)
    req = urllib.request.Request(url, headers={"User-Agent": "RAG/1.0"})
    with urllib.request.urlopen(req, timeout=30, context=ctx) as resp:
        return json.loads(resp.read())


# 获取中国历史 raw wikitext
data = api_get({
    "action": "query",
    "prop": "revisions",
    "rvprop": "content",
    "rvslots": "main",
    "titles": "中国历史",
    "format": "json",
})

pages = data["query"]["pages"]
for _pid, page in pages.items():
    raw = page["revisions"][0]["slots"]["main"]["*"]

    # 去除 wiki 标记
    text = re.sub(r"\{\{[^}]*\}\}", "", raw)
    text = re.sub(r"\[\[(?:[^|\]]*\|)?([^\]]+)\]\]", r"\1", text)
    text = re.sub(r"<ref[^>]*>.*?</ref>", "", text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"'''?", "", text)
    text = re.sub(r"={2,}\s*([^=]+?)\s*={2,}", r"\1", text)
    text = re.sub(r"\[\[(?:File|Image):[^\]]+\]\]", "", text)
    lines = [l.strip() for l in text.split("\n") if l.strip() and not l.strip().startswith("==")]
    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip()

    print(f"文章长度: {len(text)} 字符")

    doc = Document()
    doc.add_heading("中国历史", level=0)
    for para in text.split("\n"):
        stripped = para.strip()
        if stripped:
            doc.add_paragraph(stripped)

    path = "samples/中国历史.docx"
    doc.save(path)
    print(f"已保存: {path}")
    break
