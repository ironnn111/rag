#!/usr/bin/env python3
"""抓取一篇长中文维基百科文章，生成 docx。"""

import json
import os
import re
import ssl
import time
import urllib.request
import urllib.parse

try:
    from docx import Document
except ImportError:
    print("请先安装 python-docx: pip install python-docx")
    exit(1)

OUTPUT_DIR = "samples"

ctx = ssl.create_default_context()
ctx.check_hostname = False
ctx.verify_mode = ssl.CERT_NONE


def api_get(params):
    url = "https://zh.wikipedia.org/w/api.php?" + urllib.parse.urlencode(params)
    req = urllib.request.Request(url, headers={"User-Agent": "RAG/1.0"})
    with urllib.request.urlopen(req, timeout=20, context=ctx) as resp:
        return json.loads(resp.read())


# 先测几篇长文章看哪篇最长
candidates = [
    "中国历史",
    "第二次世界大战",
    "人工智能史",
    "中华民国",
    "日本历史",
    "中国",
]

best_title = ""
best_text = ""
best_len = 0

for query in candidates:
    data = api_get({"action": "query", "list": "search", "srsearch": query, "format": "json"})
    results = data["query"]["search"]
    if not results:
        continue
    best = results[0]
    title = best["title"]
    page_id = best["pageid"]

    time.sleep(0.5)

    data2 = api_get({
        "action": "query",
        "prop": "revisions",
        "rvprop": "content",
        "rvslots": "main",
        "pageids": page_id,
        "format": "json",
    })
    pages = data2["query"]["pages"]
    for _pid, page in pages.items():
        raw = page["revisions"][0]["slots"]["main"]["*"]
        text = re.sub(r"\{\{[^}]*\}\}", "", raw)
        text = re.sub(r"\[\[(?:[^|\]]*\|)?([^\]]+)\]\]", r"\1", text)
        text = re.sub(r"<ref[^>]*>.*?</ref>", "", text, flags=re.DOTALL)
        text = re.sub(r"<[^>]+>", "", text)
        text = re.sub(r"'''?", "", text)
        text = re.sub(r"={2,}\s*([^=]+?)\s*={2,}", r"\1", text)
        text = re.sub(r"\[\[(?:File|Image):[^\]]+\]\]", "", text)
        lines = [l.strip() for l in text.split("\n") if l.strip() and not l.strip().startswith("==")]
        text = "\n".join(lines)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = text.strip()
        print(f"  {title}: {len(text)} chars")
        if len(text) > best_len:
            best_len = len(text)
            best_title = title
            best_text = text
        break

    time.sleep(1)

print(f"\n选用最长篇: {best_title} ({best_len} chars)")

# 生成 docx
doc = Document()
doc.add_heading(best_title, level=0)
for para_text in best_text.split("\n"):
    stripped = para_text.strip()
    if stripped:
        doc.add_paragraph(stripped)

filename = best_title.replace("/", "_").replace("\\", "_") + ".docx"
filepath = os.path.join(OUTPUT_DIR, filename)
doc.save(filepath)
print(f"已保存: {filepath}")
